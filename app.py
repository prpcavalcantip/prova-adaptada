import streamlit as st
import fitz  # PyMuPDF
import docx
import re
from io import BytesIO
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

st.set_page_config(page_title="AdaptaProva", layout="centered")

st.title("🧠 AdaptaProva - Provas Adaptadas para Alunos com Neurodivergência")
st.markdown("Envie uma prova em PDF com texto selecionável e selecione a neurodivergência do aluno para gerar uma versão adaptada.")

# Banco de dicas para cada neurodivergência
dicas_por_tipo = {
    "TDAH": [
        "Destaque palavras-chave da pergunta.",
        "Leia a pergunta duas vezes antes de escolher a resposta.",
        "Tente eliminar as alternativas claramente erradas primeiro."
    ],
    "TEA": [
        "Preste atenção nas palavras que indicam ordem, como 'primeiro', 'depois', 'por fim'.",
        "Leia com calma. Respire fundo antes de cada pergunta.",
        "Use rascunho para organizar o que entendeu da questão."
    ],
    "Ansiedade": [
        "Lembre-se: você pode fazer uma pergunta de cada vez com calma.",
        "Respire fundo antes de começar cada questão.",
        "Você está preparado. Confie no seu raciocínio!"
    ]
}

uploaded_file = st.file_uploader("📄 Envie a prova em PDF", type=["pdf"])
tipo = st.selectbox("🧠 Neurodivergência do aluno:", ["TDAH", "TEA", "Ansiedade"])

if uploaded_file and tipo:
    st.write("✅ Arquivo carregado com sucesso. Tipo selecionado:", tipo)  # Debug
    if st.button("🔄 Gerar Prova Adaptada"):
        with st.spinner("Processando..."):

            doc = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            texto = ""
            for page in doc:
                texto += page.get_text()

            # Corrigir quebras de linha suaves
            texto = re.sub(r'(?<!\n)\n(?!\n)', ' ', texto)

            # Separar questões
            blocos = re.split(r'\bQUESTÃO\s+\d+', texto)
            blocos = [b.strip() for b in blocos if b.strip()]
            if len(blocos) > 10:
                blocos = blocos[1:]
            blocos = blocos[:10]

            docx_file = docx.Document()
            docx_file.add_heading("Prova Adaptada", 0)

            style = docx_file.styles["Normal"]
            style.font.size = Pt(14)

            # Dicas iniciais
            docx_file.add_paragraph("💡 DICAS PARA O ALUNO:", style="List Bullet")
            for dica in dicas_por_tipo[tipo]:
                p = docx_file.add_paragraph(dica, style="List Bullet")
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_after = Pt(12)
            docx_file.add_paragraph("")

            for i, bloco in enumerate(blocos):
                docx_file.add_paragraph("")  # Espaço antes da questão

                # Título da questão
                titulo = docx_file.add_paragraph()
                run = titulo.add_run(f"QUESTÃO {i+1}")
                run.bold = True
                titulo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                titulo.paragraph_format.space_after = Pt(12)

                # Separar enunciado e alternativas
                alternativas_matches = list(re.finditer(r"[A-Ea-e][\)\.].*?(?=( [A-Ea-e][\)\.]|$))", bloco, re.DOTALL))
                if alternativas_matches:
                    primeira_alternativa_pos = alternativas_matches[0].start()
                    enunciado_texto = bloco[:primeira_alternativa_pos].strip()
                    alternativas_texto = bloco[primeira_alternativa_pos:].strip()
                else:
                    enunciado_texto = bloco.strip()
                    alternativas_texto = ""

                # Enunciado
                enunciado = docx_file.add_paragraph(enunciado_texto)
                enunciado.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                enunciado.paragraph_format.line_spacing = 1.5
                enunciado.paragraph_format.space_after = Pt(24)
                docx_file.add_paragraph("")

                # Alternativas (cada uma separada por A), B), etc.)
                alternativas = re.split(r"(?=[A-Ea-e][\)\.])", alternativas_texto)
                for alt in alternativas:
                    alt = alt.strip()
                    if alt:
                        alt_par = docx_file.add_paragraph(alt)
                        alt_par.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                        alt_par.paragraph_format.line_spacing = 1.5
                        alt_par.paragraph_format.space_after = Pt(10)

                # Dicas finais
                docx_file.add_paragraph("")
                docx_file.add_paragraph("💡 Dicas para essa questão:", style="List Bullet")
                for dica in dicas_por_tipo[tipo]:
                    dica_par = docx_file.add_paragraph(dica, style="List Bullet")
                    dica_par.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    dica_par.paragraph_format.line_spacing = 1.5
                    dica_par.paragraph_format.space_after = Pt(12)

            buffer = BytesIO()
            docx_file.save(buffer)
            buffer.seek(0)

            st.success("✅ Prova adaptada gerada com sucesso!")
            st.download_button(
                label="📥 Baixar Prova Adaptada (.docx)",
                data=buffer,
                file_name="prova_adaptada.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
